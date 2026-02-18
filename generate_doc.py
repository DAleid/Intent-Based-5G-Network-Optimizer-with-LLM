"""Generate a properly formatted Word document for the project documentation."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Intent-Based 5G Network Optimizer with LLM', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Agentic AI for Autonomous Optimization of 5G-Advanced Radio Access Networks')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc.add_run(
    'A working prototype of an intent-driven network management system built on '
    '3GPP Release 18 principles. Natural language intents are parsed by a Groq-powered '
    'LLM agent pipeline and translated into network configurations, while the system '
    'continuously monitors KPIs and heals itself autonomously.'
).font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Architecture',
    '3. Agent Pipeline',
    '4. Dataset',
    '5. Network Topology',
    '6. Supported Intent Types',
    '7. Multi-Intent Conflict Resolution',
    '8. KPI Thresholds',
    '9. 3GPP Standards Compliance',
    '10. Technology Stack',
    '11. Local Setup',
    '12. Streamlit Cloud Deployment',
    '13. Project Structure',
    '14. Troubleshooting',
    '15. References',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'Traditional network management requires operators to manually translate business '
    'goals into low-level RAN parameters. This system eliminates that gap:'
)
add_bullet('An operator types: "Prioritize emergency communications at the hospital now"')
add_bullet('Six AI agents process the request end-to-end')
add_bullet('The network is reconfigured, monitored, and self-healed \u2014 no manual steps')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key Architectural Principle: ').bold = True
p.add_run(
    'LLM translates, rules validate, policies execute. '
    'The LLM is never given direct control over the network; it only parses intent and '
    'proposes configurations. A deterministic rule-based validator acts as the safety gate '
    'before any action is taken.'
)

# ═══════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Architecture', level=1)
doc.add_paragraph(
    'The system is organized into three layers:'
)

doc.add_heading('Layer 1: Streamlit UI', level=2)
add_bullet('Tab 1: Intent Processing')
add_bullet('Tab 2: Network Monitor + Topology')
add_bullet('Tab 3: Multi-Intent Conflict Resolution')

doc.add_heading('Layer 2: CrewAI Agent Pipeline (Groq LLM)', level=2)
add_bullet('Agent 1: Intent Parser (LLM)')
add_bullet('Agent 2: Reasoner (LLM)')
add_bullet('Agent 3: Validator (Rule-based) \u2190 Safety gate')
add_bullet('Agent 4: Planner (LLM)')
add_bullet('Agent 5: Monitor (LLM)')
add_bullet('Agent 6: Optimizer (LLM)')

doc.add_heading('Layer 3: Network Simulator + Real Dataset', level=2)
add_bullet('6G HetNet CSV \u2022 5,000 rows \u2022 49 Cell_IDs')
add_bullet('Dataset-driven simulator cycles through real records sequentially')

# ═══════════════════════════════════════════════════════════════
# 3. AGENT PIPELINE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Agent Pipeline', level=1)

add_table(
    ['#', 'Agent', 'Powered By', 'Role'],
    [
        ['1', 'Intent Parser', 'Groq LLM', 'Converts natural language into a structured intent object (type, priority, slice, timing, SLA targets)'],
        ['2', 'Reasoner', 'Groq LLM', 'Asks clarifying questions, assesses feasibility, identifies risks, simulates impact'],
        ['3', 'Validator', 'Rule-based', 'Hard safety gate \u2014 enforces bandwidth caps, user limits, confidence thresholds. Cannot be overridden by LLM'],
        ['4', 'Planner', 'Groq LLM', 'Generates a 3GPP-compliant network configuration (slice type, QoS, RAN params)'],
        ['5', 'Monitor', 'Groq LLM', 'Reads real KPIs from the dataset/simulator, detects anomalies, triggers alerts'],
        ['6', 'Optimizer', 'Groq LLM', 'Executes corrective actions, with automatic rollback if metrics degrade'],
    ]
)

doc.add_heading('Why Agent 3 is Rule-Based', level=2)
doc.add_paragraph(
    'The Validator enforces hard limits that must be deterministic and auditable:'
)

add_table(
    ['Limit', 'Value', 'Reason'],
    [
        ['Max bandwidth', '500 Mbps', 'Physical capacity'],
        ['Min bandwidth', '10 Mbps', 'Minimum viable allocation'],
        ['Min confidence', '70%', 'Reject uncertain LLM output'],
        ['Max cells active at once', '5', 'Prevent mass activation'],
        ['Min active cells', '5', 'Maintain coverage'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('LLMs are non-deterministic. Safety constraints must not be.')
r.italic = True

# ═══════════════════════════════════════════════════════════════
# 4. DATASET
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Dataset', level=1)

p = doc.add_paragraph()
p.add_run('File: ').bold = True
p.add_run('data/6G_HetNet_Transmission_Management.csv')

add_table(
    ['Property', 'Value'],
    [
        ['Rows', '5,000'],
        ['Columns', '24'],
        ['Unique Cell_IDs', '49 (range 1\u201349)'],
        ['Cell types', 'Macro (10), Micro (17), Pico (8), Femto (14)'],
    ]
)

doc.add_heading('Key Columns Used', level=2)
add_table(
    ['Column', 'Used For'],
    [
        ['Cell_ID', 'Pinning topology cells to specific real cells'],
        ['Cell_Type', 'Macro / Micro / Pico / Femto classification'],
        ['Achieved_Throughput_Mbps', 'Live KPI display'],
        ['Network_Latency_ms', 'SLA monitoring'],
        ['Resource_Utilization', 'Cell load %'],
        ['Signal_to_Noise_Ratio_dB', 'RAN quality metric'],
        ['Interference_Level_dB', 'Interference monitoring'],
        ['QoS_Satisfaction', 'Service quality score'],
        ['Packet_Loss_Ratio', 'Reliability metric'],
    ]
)

doc.add_heading('How the Dataset is Used', level=2)
add_bullet('Simulator reads one row per get_metrics() call, cycling through all 5,000 records sequentially', bold_prefix='Simulator: ')
add_bullet('Each of the 12 HetNet topology cells is pinned to a dedicated Cell_ID so every cell displays real, distinct KPI data from its own real-world counterpart', bold_prefix='Topology: ')

# ═══════════════════════════════════════════════════════════════
# 5. NETWORK TOPOLOGY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Network Topology', level=1)
doc.add_paragraph(
    'The UI renders a 12-cell Heterogeneous Network (HetNet) with 4 cell tiers. '
    'Each topology cell is mapped to a specific Cell_ID in the real dataset:'
)

add_table(
    ['Cell', 'Type', 'Dataset Cell_ID', 'Real Rows'],
    [
        ['C01 Central', 'Macro', '3', '103'],
        ['C02 North', 'Macro', '4', '108'],
        ['C03 South', 'Macro', '13', '98'],
        ['C04 Micro-1', 'Micro', '2', '99'],
        ['C05 Micro-2', 'Micro', '7', '78'],
        ['C06 Micro-3', 'Micro', '12', '99'],
        ['C07 Micro-4', 'Micro', '14', '94'],
        ['C08 Pico-1', 'Pico', '9', '98'],
        ['C09 Pico-2', 'Pico', '11', '102'],
        ['C10 Pico-3', 'Pico', '18', '95'],
        ['C11 Femto-1', 'Femto', '1', '102'],
        ['C12 Femto-2', 'Femto', '5', '101'],
    ]
)

doc.add_paragraph(
    'Cell health (healthy / warning / critical) is computed from real throughput, latency, and load values.'
)

# ═══════════════════════════════════════════════════════════════
# 6. SUPPORTED INTENT TYPES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Supported Intent Types', level=1)

add_table(
    ['Intent', 'Example Phrase', 'Slice'],
    [
        ['Stadium Event', '"Optimize for the match tonight"', 'eMBB'],
        ['Emergency', '"Emergency priority at the hospital"', 'URLLC'],
        ['IoT Deployment', '"10,000 sensors in the factory"', 'mMTC'],
        ['Healthcare', '"Low latency for telemedicine"', 'URLLC'],
        ['Video Streaming', '"High quality live stream"', 'eMBB'],
        ['Smart Factory', '"Industrial automation connectivity"', 'URLLC'],
        ['Gaming', '"Minimize latency for gaming"', 'URLLC'],
        ['Concert / Event', '"Social media heavy event"', 'eMBB'],
        ['Transportation', '"Vehicle connectivity on the highway"', 'URLLC'],
        ['General Optimization', '"Improve network performance"', 'eMBB'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 7. MULTI-INTENT CONFLICT RESOLUTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Multi-Intent Conflict Resolution', level=1)
doc.add_paragraph(
    'When multiple stakeholders submit competing intents simultaneously, the system resolves conflicts by:'
)
items = [
    'Ranking intents by priority (Emergency > Healthcare > Stadium > General)',
    'Allocating bandwidth from a shared 500 Mbps pool',
    'Identifying resource conflicts (overlapping cells, spectrum)',
    'Producing a negotiated configuration that satisfies as many SLAs as possible',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

# ═══════════════════════════════════════════════════════════════
# 8. KPI THRESHOLDS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. KPI Thresholds', level=1)

add_table(
    ['KPI', 'Target', 'Warning', 'Critical'],
    [
        ['Latency', '< 50 ms', '< 80 ms', '\u2265 100 ms'],
        ['Throughput', '> 100 Mbps', '> 60 Mbps', '\u2264 30 Mbps'],
        ['Packet Loss', '< 0.01%', '< 0.1%', '\u2265 1.0%'],
        ['Cell Load', '< 70%', '< 85%', '\u2265 95%'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 9. 3GPP STANDARDS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. 3GPP Standards Compliance', level=1)

add_table(
    ['Standard', 'Relevance'],
    [
        ['3GPP TS 28.312', 'Intent-Based Management'],
        ['3GPP TS 38.843', 'AI/ML for RAN'],
        ['3GPP TS 28.104', 'Self-Optimization (SON)'],
        ['3GPP Release 18', '5G-Advanced specifications'],
        ['TM Forum IG1230', 'Autonomous Networks (Level 5)'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 10. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Technology Stack', level=1)

add_table(
    ['Component', 'Technology'],
    [
        ['UI', 'Streamlit'],
        ['Agent Framework', 'CrewAI 0.86.0'],
        ['LLM', 'Groq \u2014 Llama 3.3 70B Versatile'],
        ['LLM Router', 'LiteLLM (via CrewAI)'],
        ['Visualization', 'Plotly'],
        ['Data', 'Pandas / NumPy'],
        ['Language', 'Python 3.10+'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 11. LOCAL SETUP
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Local Setup', level=1)

doc.add_heading('Prerequisites', level=2)
add_bullet('Python 3.10 or higher')
add_bullet('A free Groq API key (https://console.groq.com)')

doc.add_heading('Install', level=2)
code = doc.add_paragraph()
code.style = doc.styles['Normal']
code_text = (
    'git clone https://github.com/DAleid/Intent-Based-5G-Network-Optimizer-with-LLM.git\n'
    'cd Intent-Based-5G-Network-Optimizer-with-LLM\n'
    'python -m venv venv\n'
    'venv\\Scripts\\activate          # Windows\n'
    'source venv/bin/activate        # macOS/Linux\n'
    'pip install -r requirements.txt'
)
run = code.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Configure', level=2)
doc.add_paragraph('Copy .env.example to .env and edit it:')
p = doc.add_paragraph()
run = p.add_run('LLM_PROVIDER=groq\nGROQ_API_KEY=your_groq_api_key_here')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Run', level=2)
p = doc.add_paragraph()
run = p.add_run('streamlit run app.py')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('The app opens at http://localhost:8501.')

# ═══════════════════════════════════════════════════════════════
# 12. STREAMLIT CLOUD DEPLOYMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Streamlit Cloud Deployment', level=1)

steps = [
    'Fork or push this repository to your GitHub account.',
    'Go to share.streamlit.io and sign in with GitHub.',
    'Click New app and select this repository.',
    'Set Main file path to app.py.',
    'Under Advanced settings \u2192 Secrets, add:\n    GROQ_API_KEY = "your_groq_api_key_here"',
    'Click Deploy.',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('The dataset is included in the repository so no external data source is needed.')

# ═══════════════════════════════════════════════════════════════
# 13. PROJECT STRUCTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Project Structure', level=1)

add_table(
    ['File / Folder', 'Description'],
    [
        ['app.py', 'Streamlit UI \u2014 main entry point'],
        ['requirements.txt', 'Python dependencies'],
        ['.env.example', 'Environment variable template'],
        ['agents/crew.py', 'CrewAI crew definition + Groq LLM config'],
        ['agents/llm_client.py', 'LangChain-Groq client for tool-level LLM calls'],
        ['agents/intent_agent.py', 'Intent Parser agent definition'],
        ['agents/planner_agent.py', 'Planner agent definition'],
        ['agents/monitor_agent.py', 'Monitor agent definition'],
        ['agents/optimizer_agent.py', 'Optimizer agent definition'],
        ['tools/intent_tools.py', 'parse_intent (LLM + keyword fallback)'],
        ['tools/config_tools.py', 'generate_config, get_templates'],
        ['tools/monitor_tools.py', 'get_metrics, check_status'],
        ['tools/action_tools.py', 'execute_action, rollback'],
        ['tools/reasoning_llm.py', 'Feasibility, risk, impact reasoning'],
        ['simulator/network_sim.py', 'Dataset-driven 5G network simulator'],
        ['config/settings.py', 'LLM config, KPI thresholds, network limits'],
        ['data/6G_HetNet_Transmission_Management.csv', 'Real dataset (5,000 rows)'],
        ['data/templates.json', '3GPP slice configuration templates'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 14. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════
doc.add_heading('14. Troubleshooting', level=1)

problems = [
    ('GROQ_API_KEY not found',
     'Make sure .env exists with your key, or add it in Streamlit Cloud secrets.'),
    ('Rate limit error (429)',
     'The Groq free tier allows 12,000 tokens/minute. The app has automatic retry with '
     'exponential backoff (5s \u2192 10s \u2192 20s), but running multiple pipelines back-to-back '
     'may still trigger it. Wait 60 seconds and try again.'),
    ('Module not found',
     'Run pip install -r requirements.txt inside your virtual environment.'),
    ('Streamlit won\'t start',
     'Try: python -m streamlit run app.py'),
]
for title, fix in problems:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(fix)

# ═══════════════════════════════════════════════════════════════
# 15. REFERENCES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('15. References', level=1)

refs = [
    '3GPP TS 28.312 \u2014 Intent driven management services for mobile networks',
    '3GPP Release 18 \u2014 5G-Advanced specifications',
    'TM Forum IG1230 \u2014 Autonomous Networks Technical Architecture',
    'ETSI ZSM \u2014 Zero-touch network and Service Management',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Number')

# ── Save ──
output_path = r'c:\Users\danyh\Desktop\last trile\intent-5g-optimizer-main\Project_Documentation.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
